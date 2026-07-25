#!/usr/bin/env python3
"""
Doc Intelligence Pipeline — Core Script
=========================================
Cross-platform .docx document intelligence pipeline.

Commands:
    scan        Walk configured directories, collect .docx paths, save manifest.
    extract     Parse collected .docx files, extract text + metadata.
    knowledge   Analyze extracted text for patterns (hints, words, layout, dupes).
    actions     Build actions: generate .docx report, WhatsApp notify, or stop.
    proctor     Validate pipeline state and report health.
    full        Run scan -> extract -> knowledge -> actions -> proctor in sequence.

Usage:
    python3 pipeline.py scan
    python3 pipeline.py extract
    python3 pipeline.py knowledge
    python3 pipeline.py proctor
    python3 pipeline.py full
"""

import os
import sys
import json
import hashlib
import time
import subprocess
import re
import shutil
import glob as glob_mod
from pathlib import Path
from datetime import datetime, timezone
from collections import Counter, defaultdict
from typing import Optional

# ── Config loading ─────────────────────────────────────────────────────────
try:
    import tomli as tomllib  # Python < 3.11
except ImportError:
    try:
        import tomllib  # Python >= 3.11
    except ImportError:
        print("ERROR: tomli not installed. Run: pip install tomli")
        sys.exit(1)


def find_root() -> Path:
    """Find the pipeline root by searching upward for config.toml."""
    candidate = Path.cwd()
    for _ in range(10):
        if (candidate / "config.toml").exists():
            return candidate
        if candidate.parent == candidate:
            break
        candidate = candidate.parent
    # Fallback: platform-specific default
    if sys.platform == "win32":
        return Path("C:/DocIntel")
    return Path.home() / "DocIntel"


ROOT = find_root()
CONFIG_PATH = ROOT / "config.toml"


def load_config() -> dict:
    """Load config.toml, applying platform-specific overrides."""
    if not CONFIG_PATH.exists():
        die(f"config.toml not found at {CONFIG_PATH}. Run install.sh first.")

    with open(CONFIG_PATH, "rb") as f:
        cfg = tomllib.load(f)

    # Platform-specific scan directories
    scan = cfg.setdefault("scan", {})
    if sys.platform == "darwin":
        scan_dirs = scan.get("directories_mac", [])
    elif sys.platform == "win32":
        scan_dirs = scan.get("directories_win", [])
    else:
        scan_dirs = scan.get("directories_linux", [])

    # Expand ~ and %VAR% in paths
    expanded = []
    for d in scan_dirs:
        d = os.path.expanduser(d)
        if sys.platform == "win32":
            d = os.path.expandvars(d)
        expanded.append(d)
    scan["_resolved_dirs"] = expanded

    return cfg


def die(msg: str, code: int = 1):
    print(f"FATAL: {msg}", file=sys.stderr)
    sys.exit(code)


def log(msg: str, level: str = "INFO"):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] [{level}] {msg}")


def git_commit(message: str):
    """Stage all and commit in ROOT."""
    try:
        subprocess.run(["git", "add", "-A"], cwd=ROOT, capture_output=True, check=True)
        subprocess.run(
            ["git", "commit", "-m", message, "--allow-empty"],
            cwd=ROOT, capture_output=True, check=True,
        )
        log(f"git commit: {message}")
    except subprocess.CalledProcessError as e:
        log(f"git commit failed: {e.stderr.decode().strip()}", "WARNING")


def ensure_dirs():
    """Ensure subdirectories exist under ROOT."""
    for d in ["collected", "extracted", "knowledge", "actions", "logs"]:
        (ROOT / d).mkdir(parents=True, exist_ok=True)


# ═══════════════════════════════════════════════════════════════════════════
# SCAN: Directory-by-directory synchronous .docx collection
# ═══════════════════════════════════════════════════════════════════════════

def cmd_scan(cfg: dict):
    """Walk configured directories one at a time, collect .docx paths."""
    ensure_dirs()
    scan_cfg = cfg.get("scan", {})
    directories = scan_cfg.get("_resolved_dirs", [])
    max_depth = scan_cfg.get("max_depth", 8)
    file_types = tuple(scan_cfg.get("file_types", [".docx"]))
    min_size = scan_cfg.get("min_size_bytes", 1024)
    incremental = cfg.get("extraction", {}).get("incremental", True)

    manifest_path = ROOT / "collected" / "manifest.jsonl"

    # Load existing paths for incremental mode
    existing = set()
    if incremental and manifest_path.exists():
        with open(manifest_path, "r") as f:
            for line in f:
                line = line.strip()
                if line:
                    try:
                        existing.add(json.loads(line)["path"])
                    except (json.JSONDecodeError, KeyError):
                        pass
        log(f"Incremental mode: {len(existing)} already in manifest")

    new_count = 0
    total_size = 0

    for idx, directory in enumerate(directories, 1):
        dir_path = Path(directory).expanduser().resolve()
        if not dir_path.exists():
            log(f"SKIP [{idx}/{len(directories)}]: directory not found: {dir_path}", "WARNING")
            continue

        log(f"SCAN [{idx}/{len(directories)}]: {dir_path}")
        dir_count = 0

        # Synchronous walk — one directory at a time, no threads, no async
        for root_dir, dirs, files in os.walk(dir_path):
            # Depth limit
            depth = len(Path(root_dir).relative_to(dir_path).parts)
            if depth > max_depth:
                dirs.clear()  # don't descend further
                continue

            for fname in files:
                if not fname.lower().endswith(file_types):
                    continue

                fpath = Path(root_dir) / fname
                try:
                    fstat = fpath.stat()
                except OSError:
                    continue

                if fstat.st_size < min_size:
                    continue

                abs_path = str(fpath.resolve())

                if incremental and abs_path in existing:
                    continue

                entry = {
                    "path": abs_path,
                    "size_bytes": fstat.st_size,
                    "mtime": datetime.fromtimestamp(fstat.st_mtime, tz=timezone.utc).isoformat(),
                    "source_dir": str(dir_path),
                    "collected_at": datetime.now(timezone.utc).isoformat(),
                }

                with open(manifest_path, "a") as mf:
                    mf.write(json.dumps(entry) + "\n")

                new_count += 1
                total_size += fstat.st_size
                dir_count += 1

        log(f"  -> {dir_count} new .docx files in {dir_path}")

    log(f"Scan complete: {new_count} new files, {total_size / (1024*1024):.1f} MB total")
    git_commit(f"scan: {new_count} new .docx files collected")

    return new_count


# ═══════════════════════════════════════════════════════════════════════════
# EXTRACT: Parse .docx files, extract text + metadata
# ═══════════════════════════════════════════════════════════════════════════

def cmd_extract(cfg: dict):
    """Parse all unprocessed .docx files from the manifest."""
    ensure_dirs()
    ext_cfg = cfg.get("extraction", {})
    pause_ms = ext_cfg.get("pause_ms", 1000)
    batch_size = ext_cfg.get("batch_size", 10)

    manifest_path = ROOT / "collected" / "manifest.jsonl"
    if not manifest_path.exists():
        log("No manifest found. Run 'scan' first.", "WARNING")
        return 0

    # Load manifest
    entries = []
    with open(manifest_path, "r") as f:
        for line in f:
            line = line.strip()
            if line:
                try:
                    entries.append(json.loads(line))
                except json.JSONDecodeError:
                    pass

    # Determine which files need extraction
    extracted_dir = ROOT / "extracted"
    to_process = []
    for entry in entries:
        fhash = hashlib.md5(entry["path"].encode()).hexdigest()[:16]
        meta_file = extracted_dir / f"{fhash}.json"
        if not meta_file.exists():
            to_process.append((entry, fhash))

    if not to_process:
        log("All files already extracted. Nothing to do.")
        return 0

    log(f"Extracting {len(to_process)} files (batch_size={batch_size}, pause={pause_ms}ms)...")

    processed = 0
    failed = 0

    for i, (entry, fhash) in enumerate(to_process):
        fpath = entry["path"]
        log(f"  [{i+1}/{len(to_process)}] {Path(fpath).name}")

        try:
            # Try python-docx parsing
            try:
                from docx import Document
                doc = Document(fpath)
                text_parts = []
                heading_count = 0
                table_count = len(doc.tables)
                image_count = 0

                for para in doc.paragraphs:
                    text_parts.append(para.text)
                    if para.style and para.style.name and "Heading" in para.style.name:
                        heading_count += 1

                # Count images in relationships
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        image_count += 1

                full_text = "\n".join(text_parts)
                para_count = len(doc.paragraphs)

            except ImportError:
                log("python-docx not installed. Falling back to ZIP extraction.", "WARNING")
                # Fallback: .docx is a ZIP, extract document.xml
                import zipfile
                from xml.etree import ElementTree as ET

                with zipfile.ZipFile(fpath, "r") as zf:
                    xml_content = zf.read("word/document.xml")
                root_el = ET.fromstring(xml_content)
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                text_parts = [t.text or "" for t in root_el.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")]
                full_text = " ".join(text_parts)
                para_count = len([p for p in root_el.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")])
                heading_count = 0
                table_count = 0
                image_count = 0

            # Save extracted text
            text_path = extracted_dir / f"{fhash}.txt"
            with open(text_path, "w", encoding="utf-8") as tf:
                tf.write(full_text)

            # Save metadata
            meta = {
                "source_path": fpath,
                "hash": fhash,
                "size_bytes": entry.get("size_bytes", 0),
                "paragraphs": para_count,
                "headings": heading_count,
                "tables": table_count,
                "images": image_count,
                "char_count": len(full_text),
                "word_count": len(full_text.split()),
                "extracted_at": datetime.now(timezone.utc).isoformat(),
            }
            meta_path = extracted_dir / f"{fhash}.json"
            with open(meta_path, "w", encoding="utf-8") as mf:
                json.dump(meta, mf, indent=2)

            processed += 1

        except Exception as e:
            log(f"  FAILED: {e}", "ERROR")
            failed += 1

        # Pause between batches
        if (i + 1) % batch_size == 0 and i + 1 < len(to_process):
            log(f"  --- batch pause {pause_ms}ms ---")
            time.sleep(pause_ms / 1000.0)

    log(f"Extraction complete: {processed} processed, {failed} failed")
    git_commit(f"extract: {processed} files processed, {failed} failed")

    return processed


# ═══════════════════════════════════════════════════════════════════════════
# KNOWLEDGE: Analyze extracted text for patterns
# ═══════════════════════════════════════════════════════════════════════════

def cmd_knowledge(cfg: dict):
    """Build knowledge from extracted documents."""
    ensure_dirs()
    know_cfg = cfg.get("knowledge", {})
    extracted_dir = ROOT / "extracted"
    knowledge_dir = ROOT / "knowledge"

    # Collect all extracted texts
    docs = []
    for meta_file in sorted(extracted_dir.glob("*.json")):
        fhash = meta_file.stem
        text_file = extracted_dir / f"{fhash}.txt"
        if not text_file.exists():
            continue
        with open(meta_file, "r") as f:
            meta = json.load(f)
        with open(text_file, "r", encoding="utf-8") as f:
            text = f.read()
        docs.append({"hash": fhash, "meta": meta, "text": text})

    if len(docs) < know_cfg.get("min_corpus_size", 3):
        log(f"Need at least {know_cfg.get('min_corpus_size', 3)} documents for analysis. Found {len(docs)}.", "WARNING")
        return 0

    log(f"Building knowledge from {len(docs)} documents...")

    results = {"hints": 0, "repetitive_words": 0, "layout_patterns": 0,
               "reused_assets": 0, "duplicates": 0}

    # ── 8a: Hints Detection ────────────────────────────────────────────
    if know_cfg.get("detect_hints", True):
        hint_patterns = [
            r"(?i)\b(tip|note|hint|important|warning|caution)\s*[:;—-]",
            r"(?i)\b(should|must|recommend|always|never|ensure|avoid)\b",
            r"(?i)\b(best practice|pro tip|key takeaway|remember that)\b",
        ]
        hints_out = knowledge_dir / "hints.jsonl"
        hint_count = 0
        with open(hints_out, "w") as hf:
            for doc in docs:
                for pattern in hint_patterns:
                    for match in re.finditer(pattern, doc["text"]):
                        # Capture surrounding context
                        start = max(0, match.start() - 40)
                        end = min(len(doc["text"]), match.end() + 120)
                        snippet = doc["text"][start:end].replace("\n", " ").strip()
                        hf.write(json.dumps({
                            "doc_hash": doc["hash"],
                            "source_path": doc["meta"]["source_path"],
                            "pattern": pattern,
                            "match": match.group(0),
                            "snippet": snippet,
                        }) + "\n")
                        hint_count += 1
        log(f"  Hints: {hint_count} found")
        results["hints"] = hint_count

    # ── 8b: Repetitive Words ───────────────────────────────────────────
    if know_cfg.get("detect_repetitive_words", True):
        # Simple TF: find words appearing in >= 3 docs with high frequency
        doc_word_sets = []
        for doc in docs:
            words = re.findall(r"\b[a-zA-Z]{4,}\b", doc["text"].lower())
            doc_word_sets.append(Counter(words))

        # Find words present in >= 3 docs with significant counts
        word_doc_count = Counter()
        word_total_freq = Counter()
        for wc in doc_word_sets:
            word_doc_count.update(wc.keys())
            word_total_freq.update(wc)

        repetitive = []
        for word, ndocs in word_doc_count.items():
            if ndocs >= 3 and word_total_freq[word] >= ndocs * 3:
                repetitive.append({
                    "word": word,
                    "doc_count": ndocs,
                    "total_frequency": word_total_freq[word],
                })

        repetitive.sort(key=lambda x: x["doc_count"], reverse=True)
        repetitive = repetitive[:100]  # top 100

        rep_out = knowledge_dir / "repetitive_words.jsonl"
        with open(rep_out, "w") as rf:
            for r in repetitive:
                rf.write(json.dumps(r) + "\n")
        log(f"  Repetitive words: {len(repetitive)} found")
        results["repetitive_words"] = len(repetitive)

    # ── 8c: Layout Patterns ────────────────────────────────────────────
    if know_cfg.get("detect_layout_patterns", True):
        layout_out = knowledge_dir / "layout_patterns.jsonl"
        layout_count = 0
        with open(layout_out, "w") as lf:
            for doc in docs:
                m = doc["meta"]
                pattern = {
                    "doc_hash": doc["hash"],
                    "source_path": m["source_path"],
                    "heading_count": m.get("headings", 0),
                    "paragraph_count": m.get("paragraphs", 0),
                    "table_count": m.get("tables", 0),
                    "image_count": m.get("images", 0),
                    "char_count": m.get("char_count", 0),
                    "word_count": m.get("word_count", 0),
                }
                # Derive ratios
                if pattern["paragraph_count"] > 0:
                    pattern["tables_per_100_paras"] = round(
                        pattern["table_count"] / pattern["paragraph_count"] * 100, 1
                    )
                    pattern["headings_per_100_paras"] = round(
                        pattern["heading_count"] / pattern["paragraph_count"] * 100, 1
                    )
                lf.write(json.dumps(pattern) + "\n")
                layout_count += 1
        log(f"  Layout patterns: {layout_count} documents analyzed")
        results["layout_patterns"] = layout_count

    # ── 8d: Close Assets Reused ────────────────────────────────────────
    if know_cfg.get("detect_reused_assets", True):
        # Hash embedded images across documents
        import zipfile
        image_hashes = defaultdict(list)  # md5 -> [(doc_hash, image_name)]

        for doc in docs:
            fpath = doc["meta"]["source_path"]
            if not os.path.exists(fpath):
                continue
            try:
                with zipfile.ZipFile(fpath, "r") as zf:
                    for name in zf.namelist():
                        if name.startswith("word/media/") and not name.endswith("/"):
                            try:
                                img_data = zf.read(name)
                                img_hash = hashlib.md5(img_data).hexdigest()
                                image_hashes[img_hash].append({
                                    "doc_hash": doc["hash"],
                                    "source_path": fpath,
                                    "image_name": name,
                                })
                            except Exception:
                                pass
            except Exception:
                pass

        # Filter to images appearing in >= 2 documents
        reused = {h: refs for h, refs in image_hashes.items() if len(refs) >= 2}

        assets_out = knowledge_dir / "reused_assets.jsonl"
        asset_count = 0
        with open(assets_out, "w") as af:
            for img_hash, refs in reused.items():
                af.write(json.dumps({
                    "image_md5": img_hash,
                    "occurrence_count": len(refs),
                    "documents": [r["source_path"] for r in refs],
                }) + "\n")
                asset_count += 1
        log(f"  Reused assets: {asset_count} images found in >= 2 docs")
        results["reused_assets"] = asset_count

    # ── 8e: Duplicates (SOFT FLAG) ─────────────────────────────────────
    if know_cfg.get("detect_duplicates", True):
        threshold = know_cfg.get("similarity_threshold", 0.85)
        dup_flag_msg = know_cfg.get("duplicate_flag", "might be a small change - investigate")
        dup_out = knowledge_dir / "duplicates.jsonl"

        # Simple MinHash approximation using word n-gram shingles
        def shingles(text: str, n: int = 3) -> set:
            words = text.lower().split()
            return {" ".join(words[i:i+n]) for i in range(len(words) - n + 1)}

        def jaccard(set_a: set, set_b: set) -> float:
            if not set_a or not set_b:
                return 0.0
            return len(set_a & set_b) / len(set_a | set_b)

        dup_count = 0
        with open(dup_out, "w") as df:
            for i in range(len(docs)):
                if len(docs[i]["text"]) < 100:
                    continue
                sh_i = shingles(docs[i]["text"])
                for j in range(i + 1, len(docs)):
                    if len(docs[j]["text"]) < 100:
                        continue
                    sh_j = shingles(docs[j]["text"])
                    sim = jaccard(sh_i, sh_j)

                    # Also check size similarity
                    size_i = docs[i]["meta"].get("size_bytes", 0)
                    size_j = docs[j]["meta"].get("size_bytes", 0)
                    size_ratio = min(size_i, size_j) / max(size_i, size_j) if max(size_i, size_j) > 0 else 1.0

                    if sim > threshold or (sim > 0.7 and size_ratio > 0.95):
                        df.write(json.dumps({
                            "flag": "soft",
                            "message": dup_flag_msg,
                            "doc_a": docs[i]["meta"]["source_path"],
                            "doc_b": docs[j]["meta"]["source_path"],
                            "similarity": round(sim, 4),
                            "size_ratio": round(size_ratio, 4),
                            "recommendation": "Review for potential merge or differentiation.",
                        }) + "\n")
                        dup_count += 1

        log(f"  Duplicates (soft-flagged): {dup_count} candidate pairs")
        results["duplicates"] = dup_count

    git_commit(f"knowledge: hints={results['hints']} rep_words={results['repetitive_words']} "
               f"layouts={results['layout_patterns']} assets={results['reused_assets']} "
               f"dupes={results['duplicates']}")

    return sum(results.values())


# ═══════════════════════════════════════════════════════════════════════════
# ACTIONS: Generate .docx report, WhatsApp notify, or stop
# ═══════════════════════════════════════════════════════════════════════════

def cmd_actions(cfg: dict):
    """Build actions based on knowledge."""
    ensure_dirs()
    act_cfg = cfg.get("actions", {})
    knowledge_dir = ROOT / "knowledge"
    actions_dir = ROOT / "actions"

    report_parts = []
    report_parts.append("=== Doc Intelligence Pipeline — Summary Report ===\n")
    report_parts.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    # Count collected documents
    manifest = ROOT / "collected" / "manifest.jsonl"
    doc_count = 0
    if manifest.exists():
        with open(manifest, "r") as f:
            doc_count = sum(1 for _ in f)
    report_parts.append(f"Documents collected: {doc_count}\n")

    # Summarize knowledge
    for kfile in ["hints.jsonl", "repetitive_words.jsonl", "layout_patterns.jsonl",
                   "reused_assets.jsonl", "duplicates.jsonl"]:
        kpath = knowledge_dir / kfile
        if kpath.exists():
            with open(kpath, "r") as f:
                count = sum(1 for _ in f)
            report_parts.append(f"  {kfile}: {count} entries")

    report_text = "\n".join(report_parts)

    # ── Action: Generate .docx ─────────────────────────────────────────
    if act_cfg.get("generate_docx", True):
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc()
            doc.add_heading("Doc Intelligence Pipeline Report", level=0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"\nDocuments collected: {doc_count}")
            doc.add_paragraph("\nKnowledge Summary:")
            for kfile in ["hints.jsonl", "repetitive_words.jsonl", "layout_patterns.jsonl",
                           "reused_assets.jsonl", "duplicates.jsonl"]:
                kpath = knowledge_dir / kfile
                if kpath.exists():
                    with open(kpath, "r") as f:
                        count = sum(1 for _ in f)
                    doc.add_paragraph(f"  - {kfile}: {count} entries", style="List Bullet")

            report_path = actions_dir / f"summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(str(report_path))
            log(f"Report generated: {report_path}")
        except ImportError:
            log("python-docx not installed. Skipping .docx report generation.", "WARNING")
            log("Report text:\n" + report_text)

    # ── Action: WhatsApp Send ──────────────────────────────────────────
    if act_cfg.get("whatsapp_send", True):
        log("WhatsApp send action: requires Hermes gateway WhatsApp configured.")
        log(f"Would send: {doc_count} documents, {report_text[:200]}...")
        # In production, this would use the Hermes messaging tool or WhatsApp API.
        # The agent should call send_message or use the gateway.

    # ── Action: Stop ───────────────────────────────────────────────────
    if act_cfg.get("stop_on_complete", False):
        log("stop_on_complete=true. Pipeline terminating.")
        return 0

    return 1


# ═══════════════════════════════════════════════════════════════════════════
# PROCTOR: Validate pipeline health
# ═══════════════════════════════════════════════════════════════════════════

def cmd_proctor(cfg: dict):
    """Validate every phase of the pipeline."""
    print("=" * 60)
    print("  PROCTOR — Pipeline Validation")
    print("=" * 60)
    print()

    errors = []
    warnings = []
    ok_count = 0
    total_checks = 0

    def check(name: str, condition: bool, err_msg: str = "", warn: bool = False):
        nonlocal ok_count, total_checks
        total_checks += 1
        if condition:
            print(f"  [PASS] {name}")
            ok_count += 1
        elif warn:
            print(f"  [WARN] {name} — {err_msg}")
            warnings.append(f"{name}: {err_msg}")
        else:
            print(f"  [FAIL] {name} — {err_msg}")
            errors.append(f"{name}: {err_msg}")

    # 1. Root exists + writable
    check("Root directory exists", ROOT.exists(), f"Missing: {ROOT}")
    check("Root directory writable", os.access(ROOT, os.W_OK), f"Not writable: {ROOT}")

    # 2. Disk space >= 2GB
    mem_cfg = cfg.get("memory", {})
    required_gb = mem_cfg.get("max_disk_gb", 2)
    try:
        usage = shutil.disk_usage(ROOT)
        free_gb = usage.free / (1024 ** 3)
        check(f"Disk space >= {required_gb}GB (free: {free_gb:.1f}GB)",
              free_gb >= required_gb,
              f"Only {free_gb:.1f}GB free, need {required_gb}GB")
    except Exception as e:
        check("Disk space check", False, str(e))

    # 3. config.toml valid
    check("config.toml exists", CONFIG_PATH.exists(), str(CONFIG_PATH))
    mem_ok = cfg.get("memory", {}).get("max_disk_gb") == 2
    check("config.toml: max_disk_gb = 2GB", mem_ok,
          f"Got {cfg.get('memory', {}).get('max_disk_gb', 'missing')}")

    # 4. Git repo
    git_dir = ROOT / ".git"
    check("Git repo initialized", git_dir.exists(), str(git_dir))

    # 5. Manifest
    manifest = ROOT / "collected" / "manifest.jsonl"
    manifest_count = 0
    if manifest.exists():
        with open(manifest, "r") as f:
            manifest_count = sum(1 for _ in f)
    check("Manifest file exists", manifest.exists(), str(manifest),
          warn=(manifest_count == 0))
    if manifest_count > 0:
        print(f"         {manifest_count} .docx files in manifest")

    # 6. Extracted files match manifest
    extracted_dir = ROOT / "extracted"
    txt_count = len(list(extracted_dir.glob("*.txt"))) if extracted_dir.exists() else 0
    if manifest_count > 0:
        check(f"Extracted files ({txt_count}) vs manifest ({manifest_count})",
              txt_count >= manifest_count * 0.8,  # 80% threshold (some may fail)
              f"Only {txt_count} extracted, expected ~{manifest_count}",
              warn=True)
    else:
        check("Extraction directory exists", extracted_dir.exists(), str(extracted_dir))

    # 7. Knowledge JSONL validity
    for kfile in ["hints.jsonl", "repetitive_words.jsonl", "layout_patterns.jsonl",
                   "reused_assets.jsonl", "duplicates.jsonl"]:
        kpath = ROOT / "knowledge" / kfile
        if kpath.exists():
            valid = True
            try:
                with open(kpath, "r") as f:
                    for i, line in enumerate(f, 1):
                        if line.strip():
                            json.loads(line)
            except json.JSONDecodeError as e:
                valid = False
            check(f"knowledge/{kfile} is valid JSONL", valid, str(e) if not valid else "")
        else:
            print(f"  [INFO] knowledge/{kfile} — not yet generated")

    # 8. Dependencies
    try:
        import docx  # noqa: F401
        check("python-docx installed", True)
    except ImportError:
        check("python-docx installed", False, "Run: pip install python-docx", warn=True)

    try:
        import tomli  # noqa: F401
        check("tomli installed", True)
    except ImportError:
        try:
            import tomllib  # noqa: F401
            check("tomllib (stdlib) available", True)
        except ImportError:
            check("tomli/tomllib installed", False, "Run: pip install tomli", warn=True)

    # ── Summary ────────────────────────────────────────────────────────
    print()
    print(f"  Results: {ok_count}/{total_checks} passed")
    if errors:
        print(f"  Errors: {len(errors)}")
    if warnings:
        print(f"  Warnings: {len(warnings)}")
    print()

    if errors:
        print("ERRORS:")
        for e in errors:
            print(f"  - {e}")
        print()

    if warnings:
        print("WARNINGS:")
        for w in warnings:
            print(f"  - {w}")
        print()

    strict = cfg.get("proctor", {}).get("strict", True)
    if errors and strict:
        sys.exit(1)

    return ok_count


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════

def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    command = sys.argv[1].lower()
    cfg = load_config()
    ensure_dirs()

    commands = {
        "scan": lambda: cmd_scan(cfg),
        "extract": lambda: cmd_extract(cfg),
        "knowledge": lambda: cmd_knowledge(cfg),
        "actions": lambda: cmd_actions(cfg),
        "proctor": lambda: cmd_proctor(cfg),
        "full": lambda: run_full(cfg),
    }

    if command not in commands:
        print(f"Unknown command: {command}")
        print(__doc__)
        sys.exit(1)

    commands[command]()


def run_full(cfg: dict):
    """Run the full pipeline in sequence."""
    log("=== FULL PIPELINE START ===")
    log("Phase 5: Scanning...")
    cmd_scan(cfg)
    log("Phase 7: Extracting...")
    cmd_extract(cfg)
    log("Phase 8: Building knowledge...")
    cmd_knowledge(cfg)
    log("Phase 9: Building actions...")
    cmd_actions(cfg)
    log("Phase 10: Proctor validation...")
    cmd_proctor(cfg)
    log("=== FULL PIPELINE COMPLETE ===")


if __name__ == "__main__":
    main()
